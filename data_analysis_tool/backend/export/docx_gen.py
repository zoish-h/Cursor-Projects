"""Generate DOCX report from analysis."""
import io
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _metrics_rows(analysis: Dict[str, Any]) -> List[Tuple[str, str]]:
    dl = analysis.get("dataset_level", {})
    metrics = dl.get("metrics", {}) or {}
    return [(k, str(v)[:200]) for k, v in metrics.items() if k != "explanation_type"]


def build_docx(analysis: Dict[str, Any]) -> Tuple[bytes, str, str]:
    doc = Document()
    doc.add_heading("Data Analysis Report", 0)
    doc.add_paragraph(f"Dataset: {analysis.get('row_count', 0)} rows, {analysis.get('column_count', 0)} columns.")
    doc.add_paragraph(f"Data kind: {analysis.get('dataset_level', {}).get('characterization', {}).get('kind', 'generic')}. "
                      f"Target column: {analysis.get('dataset_level', {}).get('characterization', {}).get('target_column') or 'N/A'}.")

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(f"Columns: {', '.join(analysis.get('columns', []))}.")

    doc.add_heading("Why these metrics?", level=1)
    for msg in analysis.get("explainability", []):
        doc.add_paragraph(msg, style="List Bullet")

    doc.add_heading("Metrics", level=1)
    rows = _metrics_rows(analysis)
    if rows:
        table = doc.add_table(rows=1 + len(rows), cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        for i, (name, val) in enumerate(rows):
            row = table.rows[i + 1].cells
            row[0].text = name
            row[1].text = val
    else:
        doc.add_paragraph("No metrics computed.")

    doc.add_heading("Column profile (sample)", level=1)
    for prof in analysis.get("data_profile", [])[:10]:
        doc.add_paragraph(f"{prof.get('column', '')} ({prof.get('type', '')}): non-null {prof.get('non_null_count', 0)}.", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "analysis_report.docx"

